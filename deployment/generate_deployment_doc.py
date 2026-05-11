from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = os.path.join(os.path.dirname(__file__), "ORBITA_AWS_Deployment_Guide.docx")
doc = Document()

# Set base style
s = doc.styles['Normal']
s.font.name = 'Calibri'
s.font.size = Pt(11)

def B(txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x6E)

def h2(txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(14)

def P(txt):
    doc.add_paragraph(txt)

def Code(txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    p.paragraph_format.left_indent = Pt(20)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ORBITA - AWS EC2 Deployment Guide")
r.bold = True
r.font.size = Pt(24)

P("This guide provides step-by-step instructions to manually set up an AWS EC2 instance, map a GoDaddy domain, configure Nginx, and deploy the ORBITA backend and frontend.")

B("1. AWS EC2 Instance Setup")
P("1. Log in to AWS Console and navigate to EC2.")
P("2. Click 'Launch Instance'.")
P("3. Choose AMI: 'Ubuntu Server 24.04 LTS (HVM)'.")
P("4. Choose Instance Type: 't3a.large' (Recommended for Weaviate and 6 microservices).")
P("5. Key Pair: Create a new key pair (.pem) and download it to your local machine.")
P("6. Network Settings: Allow SSH (22), HTTP (80), and HTTPS (443) from anywhere.")
P("7. Configure Storage: Change the default 8 GB to 30 GB (gp3).")
P("8. Launch the instance.")

B("2. Mapping GoDaddy Domain to EC2")
P("1. In your AWS EC2 Dashboard, find the 'Public IPv4 Address' of your new instance.")
P("2. Log in to your GoDaddy account and navigate to 'DNS Management' for your domain.")
P("3. Edit the 'A' Record (Host: '@') and set the value to your EC2 Public IP address.")
P("4. Optional: Create a CNAME record for 'www' pointing to '@'.")
P("5. Save changes. DNS propagation may take a few hours, but usually takes 10-15 minutes.")

B("3. SSH into the Server")
P("Open your terminal and run the following commands:")
Code("chmod 400 your-key.pem")
Code("ssh -i your-key.pem ubuntu@your-ec2-ip")

B("4. Install System Dependencies")
P("Once logged in to the EC2 instance, update the system and install necessary packages.")
Code("sudo apt-get update -y")
Code("sudo apt-get install -y python3-pip python3-venv git curl nginx")

h2("Install Docker & Docker Compose")
Code("sudo apt-get install -y docker.io docker-compose")
Code("sudo systemctl start docker")
Code("sudo systemctl enable docker")
Code("sudo usermod -aG docker ubuntu")
P("Note: Log out of SSH and log back in for the docker group permissions to take effect.")

h2("Install Node.js & PM2")
Code("curl -fsSL https://deb.nodesource.com/setup_18.x | sudo -E bash -")
Code("sudo apt-get install -y nodejs")
Code("sudo npm install -g pm2")

B("5. Configure Swap Space (Crucial for Weaviate)")
P("Since Weaviate uses a lot of memory, creating a 4GB Swap file prevents the server from crashing.")
Code("sudo fallocate -l 4G /swapfile")
Code("sudo chmod 600 /swapfile")
Code("sudo mkswap /swapfile")
Code("sudo swapon /swapfile")
Code("echo '/swapfile none swap sw 0 0' | sudo tee -a /etc/fstab")

B("6. Transfer and Set Up Deployment Files")
P("Transfer your compressed 'deployment' folder from your local machine to the EC2 instance using SCP or Git.")
Code("scp -i your-key.pem deployment.tar.gz ubuntu@your-ec2-ip:~")
P("On the EC2 server, extract it:")
Code("tar -xzvf deployment.tar.gz")
Code("cd deployment")

h2("Configure .env")
P("Edit the .env file and add your production API keys and strong JWT secrets.")
Code("nano .env")

B("7. Run Backend Deployment")
P("Execute the backend setup script to create virtual environments:")
Code("./setup_backend.sh")
P("Start the Docker databases and backend PM2 services:")
Code("./start_backend.sh")

B("8. Run Frontend Deployment")
P("Execute the frontend script to install packages and build the static assets:")
Code("./setup_frontend.sh")

B("9. Configure Nginx")
P("Nginx will serve the frontend build and reverse proxy the backend APIs.")
P("First, link your deployment directory to the web root:")
Code("sudo mkdir -p /var/www")
Code("sudo ln -sfn ~/deployment /var/www/orbita")

P("Next, copy the included nginx configuration:")
Code("sudo cp ~/deployment/nginx.conf /etc/nginx/sites-available/orbita")
Code("sudo ln -sf /etc/nginx/sites-available/orbita /etc/nginx/sites-enabled/orbita")
Code("sudo rm -f /etc/nginx/sites-enabled/default")

P("Finally, restart Nginx to apply changes:")
Code("sudo systemctl restart nginx")

B("10. Verify Deployment")
P("Your ORBITA platform should now be live at your GoDaddy domain (e.g., http://yourdomain.com).")
P("To check backend health, monitor PM2 logs:")
Code("pm2 logs")

doc.save(OUT)
print(f"Document saved: {OUT}")
